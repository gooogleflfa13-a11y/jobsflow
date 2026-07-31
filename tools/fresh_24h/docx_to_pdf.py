#!/usr/bin/env python3
"""Headless DOCX → PDF for cross-industry application packages.

Fully background—no WPS, GUI or Accessibility clicks. Supports CVs and cover
letters using:
  1) LibreOffice soffice --headless
  2) Spire.Doc fallback

Usage:
  python3 tools/fresh_24h/docx_to_pdf.py path/to/file.docx
  python3 tools/fresh_24h/docx_to_pdf.py --package-dir path/to/package
"""

from __future__ import annotations

import argparse
import hashlib
import json
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from tools.io_utils import atomic_write_json

def _conversion_stamp_path(pdf: Path) -> Path:
    return pdf.with_suffix(pdf.suffix + ".jobsflow.json")


def _source_hash(docx: Path) -> str:
    digest = hashlib.sha256()
    with Path(docx).open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_conversion_stamp(docx: Path, pdf: Path, *, engine: str) -> None:
    atomic_write_json(
        _conversion_stamp_path(Path(pdf)),
        {
            "source_sha256": _source_hash(Path(docx)),
            "engine": (engine or "auto").lower(),
            "policy": "jobsflow-one-page-v1",
        },
    )


def conversion_cache_hit(docx: Path, pdf: Path, *, engine: str) -> bool:
    stamp = _conversion_stamp_path(Path(pdf))
    if not Path(docx).exists() or not Path(pdf).exists() or not stamp.exists():
        return False
    try:
        value = json.loads(stamp.read_text(encoding="utf-8"))
        return (
            value.get("source_sha256") == _source_hash(Path(docx))
            and value.get("engine") == (engine or "auto").lower()
            and value.get("policy") == "jobsflow-one-page-v1"
        )
    except (OSError, ValueError, TypeError):
        return False


def find_soffice() -> str | None:
    candidates = [
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
        Path("/opt/homebrew/bin/soffice"),
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice.bin"),
    ]
    which = subprocess.run(["which", "soffice"], capture_output=True, text=True)
    if which.returncode == 0 and which.stdout.strip():
        candidates.insert(0, Path(which.stdout.strip()))
    for c in candidates:
        if c and c.exists():
            return str(c)
    return None


def convert_libreoffice(docx: Path, pdf: Path) -> bool:
    soffice = find_soffice()
    if not soffice:
        return False
    outdir = pdf.parent
    # User profile in /tmp avoids first-run GUI and lock conflicts
    profile = Path("/tmp/lo_pdf_profile")
    profile.mkdir(parents=True, exist_ok=True)
    env_profile = f"file://{profile}"
    subprocess.check_call(
        [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--norestore",
            f"-env:UserInstallation={env_profile}",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(outdir),
            str(docx),
        ],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    produced = outdir / (docx.stem + ".pdf")
    if produced.exists() and produced != pdf:
        produced.replace(pdf)
    return pdf.exists() and pdf.stat().st_size > 1000


def densify_docx(src: Path, dst: Path, *, is_cv: bool) -> None:
    from docx import Document
    from docx.shared import Pt, Twips

    d = Document(str(src))
    margin_factor = 0.88 if is_cv else 0.90
    font_delta = -0.5 if is_cv else -0.2
    for sec in d.sections:
        sec.top_margin = Twips(int(sec.top_margin.twips * margin_factor))
        sec.bottom_margin = Twips(int(sec.bottom_margin.twips * margin_factor))
        sec.left_margin = Twips(int(sec.left_margin.twips * margin_factor))
        sec.right_margin = Twips(int(sec.right_margin.twips * margin_factor))
    for p in d.paragraphs:
        pf = p.paragraph_format
        try:
            if pf.space_after is not None:
                pf.space_after = Pt(max(0, (pf.space_after.pt or 0) * 0.30))
            if pf.space_before is not None:
                pf.space_before = Pt(max(0, (pf.space_before.pt or 0) * 0.30))
            pf.line_spacing = 1.0
        except Exception:
            pass
        for run in p.runs:
            if run.font.size:
                try:
                    run.font.size = Pt(max(8.5, run.font.size.pt + font_delta))
                except Exception:
                    pass
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    try:
                        if pf.space_after is not None:
                            pf.space_after = Pt(max(0, (pf.space_after.pt or 0) * 0.30))
                        if pf.space_before is not None:
                            pf.space_before = Pt(max(0, (pf.space_before.pt or 0) * 0.30))
                    except Exception:
                        pass
    dst.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(dst))


def _content_bbox(page) -> tuple[float, float, float, float] | None:
    blocks = page.get_text("blocks") or []
    xs0, ys0, xs1, ys1 = [], [], [], []
    for b in blocks:
        if len(b) < 5:
            continue
        text = (b[4] or "").strip() if len(b) > 4 else ""
        if not text:
            continue
        if "Evaluation" in text or "Spire.Doc" in text:
            continue
        xs0.append(b[0])
        ys0.append(b[1])
        xs1.append(b[2])
        ys1.append(b[3])
    try:
        for d in page.get_drawings() or []:
            r = d.get("rect")
            if r is None:
                continue
            xs0.append(r.x0)
            ys0.append(r.y0)
            xs1.append(r.x1)
            ys1.append(r.y1)
    except Exception:
        pass
    if not xs0:
        return None
    return (min(xs0), min(ys0), max(xs1), max(ys1))


def strip_eval_and_fit_one_page(src_pdf: Path, dst_pdf: Path) -> int:
    import fitz

    doc = fitz.open(str(src_pdf))
    for page in doc:
        for needle in (
            "Evaluation Warning: The document was created with Spire.Doc for Python.",
            "Evaluation Warning",
            "Spire.Doc for Python",
            "created with Spire",
        ):
            for rect in page.search_for(needle):
                r = fitz.Rect(rect.x0 - 2, rect.y0 - 2, rect.x1 + 2, rect.y1 + 2)
                page.add_redact_annot(r, fill=(1, 1, 1))
        page.apply_redactions()

    if doc.page_count > 1:
        w, h = doc[0].rect.width, doc[0].rect.height
        p1_text = (doc[1].get_text("text") or "").strip()
        p1_text = (
            p1_text.replace("Evaluation Warning", "")
            .replace("Spire.Doc for Python", "")
            .strip()
        )
        new = fitz.open()
        page = new.new_page(width=w, height=h)
        if not p1_text:
            page.show_pdf_page(page.rect, doc, 0)
        else:
            blocks = doc[1].get_text("blocks")
            if blocks:
                y1_max = max(b[3] for b in blocks)
                y1_min = min(b[1] for b in blocks)
                overflow_h = max(20, y1_max - y1_min + 20)
            else:
                overflow_h = 80
            total = h + overflow_h
            scale = min(0.98, h / total)
            page.show_pdf_page(fitz.Rect(0, 0, w, h * scale), doc, 0)
            page.show_pdf_page(
                fitz.Rect(0, h * scale * 0.98, w, h),
                doc,
                1,
                clip=fitz.Rect(0, 0, w, overflow_h + 40),
            )
        tmp = Path("/tmp") / f"fit1_{src_pdf.stem}.pdf"
        new.save(str(tmp))
        new.close()
        doc.close()
        doc = fitz.open(str(tmp))

    page0 = doc[0]
    w, h = page0.rect.width, page0.rect.height
    bbox = _content_bbox(page0)
    if bbox is not None:
        x0, y0, x1, y1 = bbox
        margin = 36.0
        target = fitz.Rect(margin, margin, w - margin, h - margin)
        content_h = max(1.0, y1 - y0)
        avail_h = target.height
        empty_ratio = max(0.0, (h - y1) / h)
        if empty_ratio > 0.12 and content_h < avail_h * 0.98:
            scale = min(avail_h / content_h, 1.18)
            content_w = max(1.0, x1 - x0)
            new_w = content_w * scale
            new_h = content_h * scale
            dest_x0 = target.x0 + max(0, (target.width - new_w) / 2)
            dest = fitz.Rect(dest_x0, target.y0, dest_x0 + new_w, target.y0 + new_h)
            clip = fitz.Rect(x0 - 2, y0 - 2, x1 + 2, y1 + 2)
            out = fitz.open()
            op = out.new_page(width=w, height=h)
            op.show_pdf_page(dest, doc, 0, clip=clip)
            out.save(str(dst_pdf))
            n = out.page_count
            out.close()
            doc.close()
            return n

    doc.save(str(dst_pdf))
    n = doc.page_count
    doc.close()
    return n


def convert_spire(docx: Path, pdf: Path) -> bool:
    try:
        from spire.doc import Document as SpireDoc, FileFormat
    except ImportError:
        return False
    is_cv = "CV" in docx.name and "Cover" not in docx.name
    dense = Path("/tmp") / f"dense_{docx.stem}.docx"
    raw = Path("/tmp") / f"raw_{docx.stem}.pdf"
    densify_docx(docx, dense, is_cv=is_cv)
    doc = SpireDoc()
    doc.LoadFromFile(str(dense))
    doc.SaveToFile(str(raw), FileFormat.PDF)
    doc.Close()
    strip_eval_and_fit_one_page(raw, pdf)
    return pdf.exists() and pdf.stat().st_size > 5000


def convert(
    docx: Path,
    pdf: Path | None = None,
    *,
    engine: str = "auto",
    force: bool = False,
) -> Path:
    """Convert DOCX → PDF. engine: auto | libreoffice | spire."""
    docx = docx.resolve()
    pdf = (pdf or docx.with_suffix(".pdf")).resolve()
    engine = (engine or "auto").lower()
    if not force and conversion_cache_hit(docx, pdf, engine=engine):
        print(f"OK cached (source unchanged): {pdf}")
        return pdf

    if engine in ("auto", "libreoffice"):
        if convert_libreoffice(docx, pdf):
            write_conversion_stamp(docx, pdf, engine=engine)
            print(f"OK libreoffice (headless): {pdf}")
            return pdf
        if engine == "libreoffice":
            raise RuntimeError(
                "LibreOffice soffice not found. Install LibreOffice, or use --engine spire."
            )

    if engine in ("auto", "spire"):
        if convert_spire(docx, pdf):
            write_conversion_stamp(docx, pdf, engine=engine)
            print(f"OK spire (headless): {pdf}")
            return pdf
        if engine == "spire":
            raise RuntimeError("Spire.Doc not available (pip install spire.doc).")

    raise RuntimeError(f"No headless converter available for {docx}")


def convert_package_dir(d: Path, *, engine: str = "auto", force: bool = False) -> None:
    files = sorted(d.glob("*CV.docx")) + sorted(d.glob("*Cover Letter.docx"))
    if not files:
        files = [p for p in sorted(d.glob("*.docx")) if not p.name.startswith("~$")]
    for f in files:
        convert(f, engine=engine, force=force)


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(
        description="Headless DOCX→PDF (LibreOffice or Spire). Never launches WPS."
    )
    ap.add_argument("docx", nargs="?", type=Path)
    ap.add_argument("--package-dir", type=Path)
    ap.add_argument(
        "--engine",
        choices=("auto", "libreoffice", "spire"),
        default="auto",
        help="auto: LO if installed else Spire (default). Force with libreoffice|spire.",
    )
    ap.add_argument(
        "--force",
        action="store_true",
        help="Rebuild even when DOCX content and conversion policy are unchanged.",
    )
    args = ap.parse_args(argv)

    if args.package_dir:
        convert_package_dir(args.package_dir, engine=args.engine, force=args.force)
        return 0
    if not args.docx:
        ap.error("docx or --package-dir required")
    convert(args.docx, engine=args.engine, force=args.force)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
