#!/usr/bin/env python3
"""Validate generated core job application packages."""

import argparse
import json
from pathlib import Path
from typing import Optional, Sequence

from docx import Document
from pypdf import PdfReader


REQUIRED_FILES = (
    ("job_snapshot.md", "job_snapshot.md"),
    ("申请指南.md", "申请指南.md"),
    ("application_log.md", "application_log.md"),
    ("CV.pdf", "*_CV.pdf"),
    ("CV.docx", "*_CV.docx"),
    ("Cover_Letter.pdf", "*_Cover_Letter.pdf"),
    ("Cover_Letter.docx", "*_Cover_Letter.docx"),
)

BANNED_TERMS = (
    "results-driven",
    "proven track record",
    "leverage",
    "spearhead",
    "delve",
    "testament",
    "【",
    "】",
    "TBD",
    "TODO",
    "JD candidate",
    "3.40/4.0",
    "90% satisfactory",
)

REQUIRED_JOB_FIELDS = ("company", "role", "lane", "parent_dir", "folder_name")


class ManifestError(ValueError):
    """Raised when a manifest does not match the validator's input contract."""


def _docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    for section in document.sections:
        for area in (section.header, section.footer):
            parts.extend(paragraph.text for paragraph in area.paragraphs)
    return "\n".join(parts)


def _read_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return _docx_text(path)
    return path.read_text(encoding="utf-8")


def _has_text(text: str, expected: str) -> bool:
    normalized_text = " ".join(text.split()).casefold()
    normalized_expected = " ".join(expected.split()).casefold()
    return normalized_expected in normalized_text


def _validate_text(path: Path, company: str, role: str) -> list[str]:
    try:
        text = _read_text(path)
    except Exception as error:
        return [f"{path.name}: could not open text-bearing file ({error})"]

    errors = []
    folded_text = text.casefold()
    for term in BANNED_TERMS:
        if term.casefold() in folded_text:
            errors.append(f"{path.name}: banned wording found: {term}")
    if not _has_text(text, company):
        errors.append(f"{path.name}: missing company name: {company}")
    if not _has_text(text, role):
        errors.append(f"{path.name}: missing role name: {role}")
    return errors


def _validate_pdf(path: Path) -> list[str]:
    try:
        reader = PdfReader(path)
        has_valid_page = any(
            float(page.mediabox.width) > 0 and float(page.mediabox.height) > 0
            for page in reader.pages
        )
    except Exception as error:
        return [f"{path.name}: could not open PDF ({error})"]

    if not has_valid_page:
        return [f"{path.name}: PDF has no page with nonzero dimensions"]
    return []


def _required_paths(package_dir: Path) -> tuple[list[Path], list[str]]:
    paths = []
    errors = []
    for label, pattern in REQUIRED_FILES:
        matches = sorted(package_dir.glob(pattern))
        if not matches:
            errors.append(f"missing required file: {label}")
        else:
            paths.append(matches[0])
    return paths, errors


def validate_package(package_dir: Path, company: str, role: str) -> list[str]:
    """Return package contract violations; an empty list means validation passed.

    DOCX and one-page PDF are the product contract.  A legacy ``.tex`` source,
    when present, is still scanned for banned wording but is never required.
    """
    package_dir = Path(package_dir)
    paths, errors = _required_paths(package_dir)
    for path in paths:
        if path.suffix.lower() == ".pdf":
            errors.extend(_validate_pdf(path))
        else:
            errors.extend(_validate_text(path, company, role))
    # Keep legacy source files safe without making a LaTeX toolchain a release
    # prerequisite.  This also catches stale wording when a user keeps an old
    # source beside the current DOCX/PDF materials.
    for path in sorted(package_dir.glob("*.tex")):
        if path not in paths:
            errors.extend(_validate_text(path, company, role))
    return errors


def _validate_manifest(manifest: object) -> list[dict]:
    if not isinstance(manifest, dict):
        raise ManifestError("top-level value must be an object containing a viable list")

    viable = manifest.get("viable")
    if not isinstance(viable, list):
        raise ManifestError("viable must be a list")

    for index, job in enumerate(viable):
        if not isinstance(job, dict):
            raise ManifestError(f"viable[{index}] must be an object")
        missing = [field for field in REQUIRED_JOB_FIELDS if field not in job]
        if missing:
            raise ManifestError(
                f"viable[{index}] missing required fields: {', '.join(missing)}"
            )
        invalid = [
            field
            for field in REQUIRED_JOB_FIELDS
            if not isinstance(job[field], str) or not job[field].strip()
        ]
        if invalid:
            raise ManifestError(
                f"viable[{index}] fields must be non-empty strings: {', '.join(invalid)}"
            )
    return viable


def _parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--manifest", type=Path, required=True)
    selection = parser.add_mutually_exclusive_group(required=True)
    selection.add_argument("--lane")
    selection.add_argument("--all", action="store_true", dest="all_jobs")
    return parser.parse_args(argv)


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = _parse_args(argv)
    try:
        manifest = json.loads(args.manifest.read_text(encoding="utf-8"))
        jobs = _validate_manifest(manifest)
    except (OSError, json.JSONDecodeError, ManifestError) as error:
        print(f"manifest error: {error}")
        return 2

    selected_jobs = jobs if args.all_jobs else [job for job in jobs if job["lane"] == args.lane]
    if not selected_jobs:
        if args.all_jobs:
            print("manifest error: no packages selected for --all")
        else:
            available_lanes = ", ".join(sorted({job["lane"] for job in jobs})) or "none"
            print(
                f"manifest error: no packages selected for lane {args.lane!r}; "
                f"available lanes: {available_lanes}"
            )
        return 2

    failed = 0
    for job in selected_jobs:
        package_dir = Path(job["parent_dir"]) / job["folder_name"]
        errors = validate_package(package_dir, job["company"], job["role"])
        if errors:
            failed += 1
            print(f"FAIL {job['company']} - {job['role']} ({package_dir})")
            for error in errors:
                print(f"  - {error}")
        else:
            print(f"PASS {job['company']} - {job['role']} ({package_dir})")

    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
