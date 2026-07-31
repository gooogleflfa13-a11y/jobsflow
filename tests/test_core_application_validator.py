import json
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from tools.core_applications.validate_package import validate_package


COMPANY = "Acme"
ROLE = "Legal Counsel"
VALIDATOR_SCRIPT = "tools/core_applications/validate_package.py"


def write_valid_package(package_dir: Path) -> None:
    package_dir.mkdir(parents=True, exist_ok=True)
    stem = "Yanlong_Sun_Acme_Legal_Counsel"
    content = f"{COMPANY}\n{ROLE}\n"

    for name in (
        "job_snapshot.md",
        "申请指南.md",
        "application_log.md",
        f"{stem}_CV.tex",
        f"{stem}_Cover_Letter.tex",
    ):
        (package_dir / name).write_text(content, encoding="utf-8")

    for name in (f"{stem}_CV.docx", f"{stem}_Cover_Letter.docx"):
        document = Document()
        document.add_paragraph(content)
        document.save(package_dir / name)

    for name in (f"{stem}_CV.pdf", f"{stem}_Cover_Letter.pdf"):
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        with (package_dir / name).open("wb") as output:
            writer.write(output)


def run_cli(manifest_path: Path, lane: str) -> subprocess.CompletedProcess:
    return run_cli_with_selector(manifest_path, "--lane", lane)


def run_cli_with_selector(
    manifest_path: Path, *selector: str
) -> subprocess.CompletedProcess:
    return subprocess.run(
        [
            sys.executable,
            VALIDATOR_SCRIPT,
            "--manifest",
            str(manifest_path),
            *selector,
        ],
        capture_output=True,
        text=True,
    )


def test_validator_reports_missing_contract(tmp_path: Path):
    errors = validate_package(tmp_path, COMPANY, ROLE)

    assert len(errors) == 9
    assert any("job_snapshot.md" in error for error in errors)
    assert any("CV.pdf" in error for error in errors)
    assert any("Cover_Letter.docx" in error for error in errors)


def test_validator_accepts_complete_package(tmp_path: Path):
    write_valid_package(tmp_path)

    assert validate_package(tmp_path, COMPANY, ROLE) == []


def test_validator_rejects_every_banned_or_placeholder_term(tmp_path: Path):
    banned_terms = (
        "results-driven",
        "proven track record",
        "leverage",
        "spearhead",
        "delve",
        "testament",
        "【",
        "】",
        "TBD",
        "TODO",
        "JD candidate",
        "3.40/4.0",
        "90% satisfactory",
    )

    for term in banned_terms:
        package_dir = tmp_path / str(len(list(tmp_path.iterdir())))
        write_valid_package(package_dir)
        guide = package_dir / "申请指南.md"
        guide.write_text(
            f"{COMPANY}\n{ROLE}\n{term}\n",
            encoding="utf-8",
        )

        errors = validate_package(package_dir, COMPANY, ROLE)

        assert any(term in error for error in errors), term


def test_validator_scans_tex_and_docx_text(tmp_path: Path):
    write_valid_package(tmp_path)
    tex_path = next(tmp_path.glob("*_CV.tex"))
    tex_path.write_text(
        f"{COMPANY}\n{ROLE}\nproven track record\n",
        encoding="utf-8",
    )
    docx_path = next(tmp_path.glob("*_Cover_Letter.docx"))
    document = Document()
    document.add_paragraph(f"{COMPANY}\n{ROLE}\nTODO")
    document.save(docx_path)

    errors = validate_package(tmp_path, COMPANY, ROLE)

    assert any(tex_path.name in error and "proven track record" in error for error in errors)
    assert any(docx_path.name in error and "TODO" in error for error in errors)


def test_validator_reports_missing_company_and_role_names(tmp_path: Path):
    write_valid_package(tmp_path)
    snapshot = tmp_path / "job_snapshot.md"
    snapshot.write_text("Unrelated content", encoding="utf-8")

    errors = validate_package(tmp_path, COMPANY, ROLE)

    assert any(snapshot.name in error and COMPANY in error for error in errors)
    assert any(snapshot.name in error and ROLE in error for error in errors)


def test_validator_reports_unreadable_docx(tmp_path: Path):
    write_valid_package(tmp_path)
    docx_path = next(tmp_path.glob("*_CV.docx"))
    docx_path.write_text("not a DOCX", encoding="utf-8")

    errors = validate_package(tmp_path, COMPANY, ROLE)

    assert any(docx_path.name in error and "open" in error.lower() for error in errors)


def test_validator_rejects_zero_dimension_pdf(tmp_path: Path):
    write_valid_package(tmp_path)
    pdf_path = next(tmp_path.glob("*_CV.pdf"))
    writer = PdfWriter()
    writer.add_blank_page(width=0, height=0)
    with pdf_path.open("wb") as output:
        writer.write(output)

    errors = validate_package(tmp_path, COMPANY, ROLE)

    assert any(pdf_path.name in error and "dimension" in error.lower() for error in errors)


def test_validator_reports_unreadable_pdf(tmp_path: Path):
    write_valid_package(tmp_path)
    pdf_path = next(tmp_path.glob("*_Cover_Letter.pdf"))
    pdf_path.write_text("not a PDF", encoding="utf-8")

    errors = validate_package(tmp_path, COMPANY, ROLE)

    assert any(pdf_path.name in error and "open pdf" in error.lower() for error in errors)


def test_cli_validates_only_jobs_assigned_to_lane(tmp_path: Path):
    valid_dir = tmp_path / "valid"
    write_valid_package(valid_dir)
    manifest = {
        "viable": [
            {
                "company": COMPANY,
                "role": ROLE,
                "lane": "A",
                "parent_dir": str(tmp_path),
                "folder_name": "valid",
            },
            {
                "company": "Broken Co",
                "role": "Counsel",
                "lane": "B1",
                "parent_dir": str(tmp_path),
                "folder_name": "missing",
            },
        ]
    }
    manifest_path = tmp_path / "manifest.json"
    manifest_path.write_text(json.dumps(manifest), encoding="utf-8")

    passing = run_cli(manifest_path, "A")
    failing = run_cli(manifest_path, "B1")

    assert passing.returncode == 0, passing.stdout + passing.stderr
    assert failing.returncode != 0
    assert "Broken Co" in failing.stdout


@pytest.mark.parametrize("manifest", [[], {"viable": None}])
def test_cli_rejects_malformed_manifest_structures(tmp_path: Path, manifest):
    manifest_path = tmp_path / "manifest.json"
    manifest_path.write_text(json.dumps(manifest), encoding="utf-8")

    result = run_cli(manifest_path, "A")

    assert result.returncode == 2
    assert "manifest error" in result.stdout.lower()
    assert "viable" in result.stdout.lower()
    assert "Traceback" not in result.stderr


@pytest.mark.parametrize(
    "missing_field",
    ["company", "role", "lane", "parent_dir", "folder_name"],
)
def test_cli_rejects_manifest_jobs_missing_required_fields(
    tmp_path: Path, missing_field: str
):
    job = {
        "company": COMPANY,
        "role": ROLE,
        "lane": "A",
        "parent_dir": str(tmp_path),
        "folder_name": "package",
    }
    del job[missing_field]
    manifest_path = tmp_path / "manifest.json"
    manifest_path.write_text(json.dumps({"viable": [job]}), encoding="utf-8")

    result = run_cli(manifest_path, "A")

    assert result.returncode == 2
    assert "manifest error" in result.stdout.lower()
    assert missing_field in result.stdout
    assert "Traceback" not in result.stderr


@pytest.mark.parametrize("lane", ["", "UNKNOWN"])
def test_cli_rejects_lane_with_zero_selected_packages(tmp_path: Path, lane: str):
    manifest = {
        "viable": [
            {
                "company": COMPANY,
                "role": ROLE,
                "lane": "A",
                "parent_dir": str(tmp_path),
                "folder_name": "package",
            }
        ]
    }
    manifest_path = tmp_path / "manifest.json"
    manifest_path.write_text(json.dumps(manifest), encoding="utf-8")

    result = run_cli(manifest_path, lane)

    assert result.returncode == 2
    assert "no packages selected" in result.stdout.lower()
    assert repr(lane) in result.stdout
    assert "Traceback" not in result.stderr


def test_cli_all_validates_every_viable_job(tmp_path: Path):
    valid_dir = tmp_path / "valid"
    write_valid_package(valid_dir)
    manifest = {
        "viable": [
            {
                "company": COMPANY,
                "role": ROLE,
                "lane": "A",
                "parent_dir": str(tmp_path),
                "folder_name": "valid",
            },
            {
                "company": "Broken Co",
                "role": "Counsel",
                "lane": "B1",
                "parent_dir": str(tmp_path),
                "folder_name": "missing",
            },
        ]
    }
    manifest_path = tmp_path / "manifest.json"
    manifest_path.write_text(json.dumps(manifest), encoding="utf-8")

    result = run_cli_with_selector(manifest_path, "--all")

    assert result.returncode == 1
    assert f"PASS {COMPANY}" in result.stdout
    assert "FAIL Broken Co" in result.stdout


@pytest.mark.parametrize("selector", [(), ("--lane", "A", "--all")])
def test_cli_requires_exactly_one_selection_mode(tmp_path: Path, selector):
    manifest_path = tmp_path / "manifest.json"
    manifest_path.write_text(json.dumps({"viable": []}), encoding="utf-8")

    result = run_cli_with_selector(manifest_path, *selector)

    assert result.returncode == 2
    assert "--lane" in result.stderr
    assert "--all" in result.stderr


def test_cli_all_rejects_zero_selected_packages(tmp_path: Path):
    manifest_path = tmp_path / "manifest.json"
    manifest_path.write_text(json.dumps({"viable": []}), encoding="utf-8")

    result = run_cli_with_selector(manifest_path, "--all")

    assert result.returncode == 2
    assert "no packages selected" in result.stdout.lower()
    assert "--all" in result.stdout
